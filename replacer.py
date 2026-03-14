"""
replacer.py — Replace {DOCID} markers in a .docx with tracked changes.

Marker format inside Word documents: {XXXXX}  e.g. {AB123}

First apply:
  Finds {XXXXX}, replaces with:
    [hidden run: {XXXXX}]  [w:ins: reference]

Re-apply:
  Finds hidden {XXXXX} run, reads effective visible text after it,
  replaces with:
    [hidden run: {XXXXX}]  [w:del: old_text]  [w:ins: new_reference]

Every change is a proper Word tracked change — visible in the revision pane,
accept/rejectable in Word.

Dependencies:
    pip install python-docx
"""

import copy
from datetime import datetime, timezone
from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

XML_SPACE = "{http://www.w3.org/XML/1998/namespace}space"
TC_AUTHOR  = "Crossref"


# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------

def _now_iso() -> str:
    return datetime.now(timezone.utc).strftime("%Y-%m-%dT%H:%M:%SZ")


def _iter_paragraphs(doc):
    yield from doc.paragraphs
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                yield from cell.paragraphs


def _has_vanish(run_elem) -> bool:
    rpr = run_elem.find(qn("w:rPr"))
    return rpr is not None and rpr.find(qn("w:vanish")) is not None


def _run_text(run_elem) -> str:
    return "".join(t.text or "" for t in run_elem.iter(qn("w:t")))


def _set_run_text(run_elem, text: str):
    t_elems = list(run_elem.iter(qn("w:t")))
    if t_elems:
        t_elems[0].text = text
        t_elems[0].set(XML_SPACE, "preserve")
        for extra in t_elems[1:]:
            run_elem.remove(extra)


def _run_text_map(runs):
    full, pos_map = [], []
    for ri, run in enumerate(runs):
        t = run.text or ""
        full.append(t)
        for ci in range(len(t)):
            pos_map.append((ri, ci))
    return "".join(full), pos_map


def _next_rev_id(doc) -> int:
    max_id = 0
    for tag in (qn("w:ins"), qn("w:del")):
        for elem in doc.element.iter(tag):
            try:
                max_id = max(max_id, int(elem.get(qn("w:id"), 0)))
            except (ValueError, TypeError):
                pass
    return max_id + 1


# ---------------------------------------------------------------------------
# Run / element factories
# ---------------------------------------------------------------------------

def _make_hidden_run(text: str):
    """Invisible run — stores the original {XXXXX} marker permanently."""
    run = OxmlElement("w:r")
    rpr = OxmlElement("w:rPr")
    rpr.append(OxmlElement("w:vanish"))
    run.append(rpr)
    t = OxmlElement("w:t")
    t.text = text
    t.set(XML_SPACE, "preserve")
    run.append(t)
    return run


def _make_visible_run(source_run_elem, text: str):
    """Clone a run's character formatting (minus vanish) with new text."""
    run = copy.deepcopy(source_run_elem)
    rpr = run.find(qn("w:rPr"))
    if rpr is not None:
        for v in rpr.findall(qn("w:vanish")):
            rpr.remove(v)
    _set_run_text(run, text)
    return run


def _make_ins(text: str, rev_id: int, date: str, source_run=None):
    """<w:ins> tracked insertion."""
    ins = OxmlElement("w:ins")
    ins.set(qn("w:id"), str(rev_id))
    ins.set(qn("w:author"), TC_AUTHOR)
    ins.set(qn("w:date"), date)
    run = _make_visible_run(source_run, text) if source_run is not None else _plain_run(text)
    ins.append(run)
    return ins


def _make_del(text: str, rev_id: int, date: str, source_run=None):
    """<w:del> tracked deletion (uses w:delText inside run)."""
    del_elem = OxmlElement("w:del")
    del_elem.set(qn("w:id"), str(rev_id))
    del_elem.set(qn("w:author"), TC_AUTHOR)
    del_elem.set(qn("w:date"), date)

    if source_run is not None:
        run = copy.deepcopy(source_run)
        rpr = run.find(qn("w:rPr"))
        if rpr is not None:
            for v in rpr.findall(qn("w:vanish")):
                rpr.remove(v)
        # Replace all w:t with w:delText
        for t_elem in list(run.iter(qn("w:t"))):
            parent = t_elem.getparent()
            del_text = OxmlElement("w:delText")
            del_text.text = text
            del_text.set(XML_SPACE, "preserve")
            parent.replace(t_elem, del_text)
            break
        for extra in list(run.iter(qn("w:t"))):
            extra.getparent().remove(extra)
    else:
        run = OxmlElement("w:r")
        del_text = OxmlElement("w:delText")
        del_text.text = text
        del_text.set(XML_SPACE, "preserve")
        run.append(del_text)

    del_elem.append(run)
    return del_elem


def _plain_run(text: str):
    run = OxmlElement("w:r")
    t = OxmlElement("w:t")
    t.text = text
    t.set(XML_SPACE, "preserve")
    run.append(t)
    return run


# ---------------------------------------------------------------------------
# Find the hidden marker and what follows it
# ---------------------------------------------------------------------------

def _find_hidden_marker(doc, marker: str):
    """
    Find hidden run with text == marker.
    Returns (p_element, hidden_run_elem) or (None, None).
    """
    for para in _iter_paragraphs(doc):
        p = para._p
        for elem in p:
            if elem.tag == qn("w:r") and _has_vanish(elem) and _run_text(elem) == marker:
                return p, elem
    return None, None


def _collect_after_hidden(p, hidden_elem):
    """
    Collect w:r / w:ins / w:del elements immediately following hidden_elem.
    Returns (effective_visible_text, elements_to_remove, source_run_for_formatting).
    effective_visible_text is from w:ins or w:r (not w:del, which is already deleted).
    """
    children = list(p)
    idx = children.index(hidden_elem) + 1
    collected = []
    effective_text = ""
    source_run = None

    # Tags that are non-content decorations — skip over them without breaking
    _SKIP_TAGS = {qn("w:bookmarkStart"), qn("w:bookmarkEnd"), qn("w:proofErr"), qn("w:rPrChange")}

    for elem in children[idx:]:
        tag = elem.tag
        if tag == qn("w:r"):
            effective_text = _run_text(elem)
            source_run = elem
            collected.append(elem)
            break
        elif tag == qn("w:ins"):
            for r in elem:
                if r.tag == qn("w:r"):
                    effective_text = _run_text(r)
                    source_run = r
                    break
            collected.append(elem)
            break
        elif tag == qn("w:del"):
            collected.append(elem)
            # keep going — a w:ins may follow
        elif tag in _SKIP_TAGS:
            continue  # skip non-content decorations
        else:
            break

    return effective_text, collected, source_run


# ---------------------------------------------------------------------------
# First replacement: find marker text, insert hidden + tracked ins
# ---------------------------------------------------------------------------

def _first_replacement(doc, marker: str, reference: str) -> bool:
    for para in _iter_paragraphs(doc):
        runs = para.runs
        if not runs:
            continue

        full_text, pos_map = _run_text_map(runs)
        idx = full_text.find(marker)
        if idx == -1:
            continue

        end_idx = idx + len(marker) - 1
        if end_idx >= len(pos_map):
            continue

        start_run_i, start_char_i = pos_map[idx]
        end_run_i,   end_char_i   = pos_map[end_idx]

        p = para._p
        run_elems   = [r._r for r in runs]
        source_run  = run_elems[start_run_i]
        before_text = runs[start_run_i].text[:start_char_i]
        after_text  = runs[end_run_i].text[end_char_i + 1:]

        # Remove merged runs
        for ri in range(start_run_i + 1, end_run_i + 1):
            p.remove(run_elems[ri])

        # Insert before_run if needed
        if before_text:
            p.insert(list(p).index(source_run), _make_visible_run(source_run, before_text))

        source_idx  = list(p).index(source_run)
        rev_id      = _next_rev_id(doc)
        date        = _now_iso()
        hidden_run  = _make_hidden_run(marker)
        ins_elem    = _make_ins(reference, rev_id, date, source_run)

        p.remove(source_run)
        p.insert(source_idx, ins_elem)
        p.insert(source_idx, hidden_run)

        # Insert after_run if needed
        if after_text:
            after_pos = list(p).index(ins_elem) + 1
            p.insert(after_pos, _make_visible_run(source_run, after_text))

        return True

    return False


# ---------------------------------------------------------------------------
# Public API
# ---------------------------------------------------------------------------

def apply_reference(filepath: str, doc_id: str, reference: str) -> dict:
    """
    Apply `reference` to the .docx at `filepath` as a tracked change.

    Returns:
        { "old_text": str, "new_text": str, "method": "search"|"hidden_marker" }

    Raises ValueError for unsupported file types or if marker not found.
    """
    if not filepath.lower().endswith(".docx"):
        raise ValueError(f"Only .docx files are supported (got: {filepath})")

    reference = reference.strip()
    marker = f"{{{doc_id}}}"
    doc    = Document(filepath)

    p, hidden_run = _find_hidden_marker(doc, marker)

    if hidden_run is not None:
        # Re-apply: replace whatever follows the hidden marker with del+ins
        effective_text, to_remove, source_run = _collect_after_hidden(p, hidden_run)
        effective_text = effective_text.strip()

        rev_id  = _next_rev_id(doc)
        date    = _now_iso()

        # Remove old tracked content
        for elem in to_remove:
            p.remove(elem)

        # Insert after hidden run: [del: old][ins: new]
        insert_pos = list(p).index(hidden_run) + 1
        if effective_text:
            p.insert(insert_pos, _make_del(effective_text, rev_id,     date, source_run))
            p.insert(insert_pos + 1, _make_ins(reference,  rev_id + 1, date, source_run))
        else:
            p.insert(insert_pos, _make_ins(reference, rev_id, date, source_run))

        old_text = effective_text
        method   = "hidden_marker"
    else:
        # First apply: find marker text and replace
        if not _first_replacement(doc, marker, reference):
            raise ValueError(f"Could not find '{marker}' in document")
        old_text = marker
        method   = "search"

    doc.save(filepath)
    return {"old_text": old_text, "new_text": reference, "method": method}
